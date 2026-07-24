"""Word 生成：月考打字截图文档"""
import io
import base64
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_cn_font(run, font_name='微软雅黑', size=None):
    """设置中西文字体"""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = size


def generate_exam_word(records, month_str, images=None):
    """生成月考打字截图 Word 文档"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    # 设置 Normal 样式默认中西文字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.paragraph_format.space_after = Pt(4)

    if images is None:
        images = {}

    for i, rec in enumerate(records):
        name = rec.get("name", "")
        score = rec.get("score", 0)

        # 姓名段落
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_name = p_name.add_run(f"姓名：{name}")
        set_cn_font(run_name, '微软雅黑', Pt(12))

        # 成绩段落
        p_score = doc.add_paragraph()
        p_score.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_score = p_score.add_run(f"打字字数：{score}字/分")
        set_cn_font(run_score, '微软雅黑', Pt(12))

        # 截图（如果有）
        img_data = images.get(str(i)) or images.get(name)
        if img_data:
            try:
                img_bytes = base64.b64decode(img_data)
                img_stream = io.BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Inches(5.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            except Exception:
                pass

        # 每两个人之间加分隔（最后一人不加）
        if i < len(records) - 1:
            doc.add_paragraph()

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
